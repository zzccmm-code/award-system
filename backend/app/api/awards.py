import os, re as re_mod, tempfile, threading, logging
from datetime import datetime
import docx
from sqlalchemy import case
from fastapi import APIRouter, Depends, Query, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_
from typing import Optional
logger = logging.getLogger(__name__)

from ..database import get_db, SessionLocal
from ..models import Award
from ..schemas import AwardOut, AwardListResponse, AwardCreate
from ..excel_handler import generate_template, parse_excel, generate_export
from ..pdf_parser import parse_pdf, _parse_award_text
from ..keywords import is_sichuan_related
from .. import progress as prog_store
from ..auth import get_current_user

router = APIRouter(prefix="/api/awards", tags=["awards"], dependencies=[Depends(get_current_user)])

def _clean_text(val):
    if isinstance(val, str): return re_mod.sub(r"\s+", "", val)
    return val

def _save_entries(entries: list[dict], db: Session, page_title: str = "", skip_filter: bool = False) -> tuple[int, int]:
    saved, skipped = 0, 0
    for entry in entries:
        entry["project_name"] = _clean_text(entry["project_name"])
        entry["completing_unit"] = _clean_text(entry["completing_unit"])
        entry["completers"] = _clean_text(entry["completers"])
        if entry.get("source","").startswith("PDF") and page_title: entry["source"] = page_title[:100]
        is_rel, mkw = (True, []) if skip_filter else is_sichuan_related(entry.get("completing_unit", ""))
        if not is_rel: skipped += 1; continue
        existing = db.query(Award).filter(Award.project_name == entry["project_name"], Award.award_year == entry["award_year"], Award.source == entry["source"]).first()
        if not existing:
            entry["matched_keywords"] = ",".join(mkw); entry["is_power_related"] = 1
            db.add(Award(**entry)); saved += 1
    db.commit()
    return saved, skipped

def _parse_docx_table(doc):
    # Detect source from document title
    source = "Word\u6587\u6863\u5bfc\u5165"
    for p in doc.paragraphs[:5]:
        text = p.text.strip()
        for kw in ["\u6cb3\u5317\u7701","\u5c71\u4e1c\u7701","\u5e7f\u4e1c\u7701","\u6d59\u6c5f\u7701","\u6c5f\u82cf\u7701","\u56db\u5ddd\u7701","\u6cb3\u5357\u7701","\u6e56\u5317\u7701","\u6e56\u5357\u7701","\u798f\u5efa\u7701","\u5b89\u5fbd\u7701","\u6c5f\u897f\u7701","\u8fbd\u5b81\u7701","\u5409\u6797\u7701","\u9ed1\u9f99\u6c5f\u7701","\u9655\u897f\u7701","\u7518\u8083\u7701","\u9752\u6d77\u7701","\u4e91\u5357\u7701","\u8d35\u5dde\u7701","\u6d77\u5357\u7701","\u5c71\u897f\u7701","\u897f\u85cf\u81ea\u6cbb\u533a","\u65b0\u7586\u7ef4\u543e\u5c14\u81ea\u6cbb\u533a","\u5185\u8499\u53e4\u81ea\u6cbb\u533a","\u5e7f\u897f\u58ee\u65cf\u81ea\u6cbb\u533a","\u5b81\u590f\u56de\u65cf\u81ea\u6cbb\u533a","\u5317\u4eac\u5e02","\u4e0a\u6d77\u5e02","\u5929\u6d25\u5e02","\u91cd\u5e86\u5e02"]:
            if kw in text: source = kw; break
        if source != "Word\u6587\u6863\u5bfc\u5165": break
    for table in doc.tables:
        if len(table.rows) < 2: continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        hdr_text = " ".join(headers)
        if "项目名称" not in hdr_text: continue
        col_map = {}
        for i, h in enumerate(headers):
            if "项目名称" in h: col_map["name"] = i
            elif "完成单位" in h: col_map["unit"] = i
            elif "完成人" in h: col_map["people"] = i
            elif "等级" in h: col_map["level"] = i
            elif "类别" in h: col_map["type"] = i
        if "name" not in col_map: continue
        entries = []
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            name = cells[col_map.get("name", -1)] if col_map.get("name", -1) >= 0 and col_map.get("name", -1) < len(cells) else ""
            if not name or len(name) < 4: continue
            entries.append({
                "project_name": _clean_text(name),
                "award_year": datetime.now().year,
                "award_type": cells[col_map.get("type", -1)] if col_map.get("type", -1) >= 0 and col_map.get("type", -1) < len(cells) else "科技进步奖",
                "award_level": cells[col_map.get("level", -1)] if col_map.get("level", -1) >= 0 and col_map.get("level", -1) < len(cells) else "",
                "completing_unit": _clean_text(cells[col_map.get("unit", -1)]) if col_map.get("unit", -1) >= 0 and col_map.get("unit", -1) < len(cells) else "",
                "completers": _clean_text(cells[col_map.get("people", -1)]) if col_map.get("people", -1) >= 0 and col_map.get("people", -1) < len(cells) else "",
                "source": source,
            })
        if entries: return entries
    return None

@router.get("", response_model=AwardListResponse)
def search_awards(keyword: Optional[str] = Query(None), award_year: Optional[int] = Query(None), award_type: Optional[str] = Query(None), award_category: Optional[str] = Query(None), page: int = Query(1, ge=1), page_size: int = Query(10, ge=1, le=100), db: Session = Depends(get_db)):
    query = db.query(Award)
    if keyword:
        kw = f"%{keyword}%"
        query = query.filter(or_(Award.project_name.ilike(kw), Award.completing_unit.ilike(kw), Award.completers.ilike(kw), Award.source.ilike(kw)))
    if award_year: query = query.filter(Award.award_year == award_year)
    if award_type: query = query.filter(Award.award_type == award_type)
    if award_category: query = query.filter(Award.award_category == award_category)
    total = query.count()
    items = query.order_by(
        case(
            (Award.award_category == "\u56fd\u5bb6\u7ea7", 1),
            (Award.award_category == "\u7701\u90e8\u7ea7", 2),
            (Award.award_category == "\u884c\u4e1a\u7ea7", 3),
            (Award.award_category == "\u56fd\u7f51\u516c\u53f8\u7ea7", 4),
            (Award.award_category == "\u7701\u516c\u53f8\u7ea7", 5),
            else_=6,
        ),
        Award.award_year.desc(),
        case(
            (Award.award_level == "\u7279\u7b49\u5956", 1),
            (Award.award_level == "\u4e00\u7b49\u5956", 2),
            (Award.award_level == "\u4e8c\u7b49\u5956", 3),
            (Award.award_level == "\u4e09\u7b49\u5956", 4),
            else_=5,
        ),
    ).offset((page-1)*page_size).limit(page_size).all()
    return AwardListResponse(total=total, items=[AwardOut.model_validate(item) for item in items], page=page, page_size=page_size)

@router.get("/categories", response_model=list[str])
def get_award_categories(db: Session = Depends(get_db)):
    categories = db.query(Award.award_category).distinct().all()
    return sorted([c[0] for c in categories if c[0]])


@router.get("/types", response_model=list[str])
def get_award_types(db: Session = Depends(get_db)):
    types = db.query(Award.award_type).distinct().all()
    return sorted([t[0] for t in types if t[0]])

@router.get("/years", response_model=list[int])
def get_award_years(db: Session = Depends(get_db)):
    years = db.query(Award.award_year).distinct().all()
    return sorted([y[0] for y in years], reverse=True)

@router.get("/template/download")
def download_template():
    buffer = generate_template()
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=award_import_template.xlsx"})

@router.post("/upload/excel")
def upload_excel(file: UploadFile = File(...), db: Session = Depends(get_db)):
    if not file.filename or not file.filename.endswith((".xlsx",".xls")): raise HTTPException(400, "??? .xlsx ? .xls ??")
    contents = file.file.read()
    try: entries = parse_excel(contents)
    except Exception as e: raise HTTPException(400, f"??Excel??: {str(e)}")
    if not entries: raise HTTPException(400, "Excel????????")
    saved, skipped = _save_entries(entries, db, skip_filter=True)
    return {"status":"ok","total": len(entries),"saved":saved,"skipped":skipped}

@router.post("/upload/file")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename: raise HTTPException(400, "No file")
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in (".pdf", ".docx"): raise HTTPException(400, "??? PDF ? Word (.docx) ??")
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    tmp.write(file.file.read()); tmp_path = tmp.name; tmp.close()
    task_id = prog_store.create_task()
    def worker():
        dbi = SessionLocal()
        try:
            prog_store.update(task_id, 0, 1, "Processing...")
            entries = []
            if ext == ".docx":
                doc = docx.Document(tmp_path)
                doc_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells: doc_text += "\n" + " | ".join(cells)
                entries = _parse_docx_table(doc) or _parse_award_text(doc_text)
                prog_store.update(task_id, 1, 1, "Parsed Word document")
            else:
                def on_progress(cur, tot):
                    prog_store.update(task_id, cur, tot, f"OCR: {cur}/{tot}")
                entries = parse_pdf(tmp_path, progress_callback=on_progress)
            os.unlink(tmp_path)
            if not entries:
                prog_store.fail(task_id, "???????????"); return
            saved, skipped = _save_entries(entries, dbi)
            prog_store.complete(task_id, {"total": len(entries), "saved": saved, "skipped": skipped})
        except Exception as e:
            prog_store.fail(task_id, str(e))
        finally:
            dbi.close()
    threading.Thread(target=worker, daemon=True).start()
    return {"status":"processing","task_id":task_id}

@router.put("/{award_id}")
def update_award(award_id: int, data: dict, db: Session = Depends(get_db)):
    award = db.query(Award).filter(Award.id == award_id).first()
    if not award: raise HTTPException(404, "Not found")
    for field in ["project_name","award_year","award_type","award_level","award_category","completing_unit","completers","source"]:
        if field in data: setattr(award, field, _clean_text(data[field]) if field in ("project_name","completing_unit","completers") else data[field])
    db.commit(); return {"status":"ok"}

@router.delete("/{award_id}")
def delete_award(award_id: int, db: Session = Depends(get_db)):
    award = db.query(Award).filter(Award.id == award_id).first()
    if not award: raise HTTPException(404, "Not found")
    db.delete(award); db.commit(); return {"status":"ok"}

@router.post("/export")
def export_awards(data: dict, db: Session = Depends(get_db)):
    ids = data.get("ids", []) if data else []
    query = db.query(Award)
    if ids: query = query.filter(Award.id.in_(ids))
    items = query.order_by(Award.award_year.desc()).all()
    rows = [{"id": a.id, "project_name": a.project_name, "award_year": a.award_year, "award_type": a.award_type, "award_level": a.award_level, "completing_unit": a.completing_unit, "completers": a.completers, "source": a.source} for a in items]
    buffer = generate_export(rows)
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=awards_export.xlsx"})

@router.get("/upload/pdf/progress/{task_id}")
def get_pdf_progress(task_id: str):
    return prog_store.get(task_id)
